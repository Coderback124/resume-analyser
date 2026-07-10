from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from io import BytesIO
from pathlib import Path
from typing import Optional
import logging
import re

from docx import Document
from PyPDF2 import PdfReader


router = APIRouter()
logger = logging.getLogger(__name__)


# =========================================================
# CONFIGURATION
# =========================================================

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "application/octet-stream",
}

MAX_FILE_SIZE = 5 * 1024 * 1024
MAX_SCORE = 100
MAX_JOB_DESCRIPTION_LENGTH = 50_000
MIN_EXTRACTED_TEXT_LENGTH = 100


# =========================================================
# TEXT CONSTANTS
# =========================================================

SECTION_GROUPS = {
    "contact": [
        "email",
        "phone",
        "linkedin",
        "github",
    ],
    "summary": [
        "summary",
        "profile",
        "professional summary",
        "professional profile",
        "objective",
        "about me",
    ],
    "experience": [
        "experience",
        "work experience",
        "employment",
        "employment history",
        "professional experience",
        "work history",
    ],
    "education": [
        "education",
        "academic background",
        "qualifications",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "technologies",
        "tech stack",
    ],
    "projects": [
        "projects",
        "personal projects",
        "academic projects",
        "portfolio",
    ],
}


TECHNICAL_SKILLS = {
    "python",
    "java",
    "javascript",
    "typescript",
    "c++",
    "c#",
    "sql",
    "html",
    "css",
    "react",
    "react native",
    "node.js",
    "nodejs",
    "fastapi",
    "django",
    "flask",
    "spring",
    "spring boot",
    "postgresql",
    "mysql",
    "sqlite",
    "mongodb",
    "redis",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "git",
    "github",
    "linux",
    "ubuntu",
    "rest api",
    "restful api",
    "api",
    "sqlalchemy",
    "pandas",
    "numpy",
    "machine learning",
    "artificial intelligence",
    "ai",
    "data analysis",
    "power bi",
    "tableau",
    "agile",
    "scrum",
    "ci/cd",
    "devops",
    "cloud",
    "testing",
    "unit testing",
    "automation",
    "cybersecurity",
    "networking",
}


SOFT_SKILLS = {
    "communication",
    "leadership",
    "teamwork",
    "collaboration",
    "problem solving",
    "problem-solving",
    "analytical",
    "adaptability",
    "time management",
    "attention to detail",
    "critical thinking",
}


ACTION_VERBS = {
    "achieved",
    "analysed",
    "analyzed",
    "automated",
    "built",
    "collaborated",
    "created",
    "delivered",
    "designed",
    "developed",
    "engineered",
    "implemented",
    "improved",
    "increased",
    "integrated",
    "led",
    "managed",
    "optimized",
    "reduced",
    "resolved",
    "supported",
    "tested",
}


STOP_WORDS = {
    "a", "an", "and", "are", "as", "at", "be", "been", "being",
    "but", "by", "can", "for", "from", "has", "have", "having",
    "he", "her", "here", "hers", "him", "his", "how", "i", "if",
    "in", "into", "is", "it", "its", "may", "more", "most", "must",
    "of", "on", "or", "our", "ours", "should", "so", "such", "than",
    "that", "the", "their", "theirs", "them", "then", "there", "these",
    "they", "this", "those", "to", "too", "us", "use", "using", "very",
    "was", "we", "were", "what", "when", "where", "which", "who",
    "will", "with", "would", "you", "your", "yours",
    "job", "role", "position", "candidate", "company", "work",
    "working", "required", "preferred", "responsibilities",
    "responsibility", "requirements", "qualification", "qualifications",
    "including", "include", "includes", "team", "opportunity",
}


# =========================================================
# TEXT HELPERS
# =========================================================

def normalise_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def contains_term(text: str, term: str) -> bool:
    pattern = rf"(?<!\w){re.escape(term.lower())}(?!\w)"
    return re.search(pattern, text.lower()) is not None


def contains_any(text: str, terms) -> bool:
    return any(contains_term(text, term) for term in terms)


def clamp_score(value: float) -> int:
    return max(0, min(100, round(value)))


def extract_words(text: str) -> list[str]:
    return re.findall(
        r"[a-zA-Z][a-zA-Z0-9+#./-]{1,}",
        text.lower(),
    )


def extract_meaningful_keywords(text: str) -> set[str]:
    words = extract_words(text)

    return {
        word
        for word in words
        if word not in STOP_WORDS
        and len(word) >= 3
        and not word.isdigit()
    }


def find_known_skills(text: str) -> set[str]:
    text_lower = text.lower()

    all_skills = TECHNICAL_SKILLS | SOFT_SKILLS

    return {
        skill
        for skill in all_skills
        if contains_term(text_lower, skill)
    }


def find_sections(text: str) -> dict[str, bool]:
    text_lower = text.lower()

    return {
        section: contains_any(text_lower, aliases)
        for section, aliases in SECTION_GROUPS.items()
    }


# =========================================================
# FILE EXTRACTION
# =========================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))

        if len(reader.pages) > 20:
            raise ValueError("PDF contains too many pages")

        pages = []

        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)

        return "\n".join(pages)

    except Exception as exc:
        logger.exception("PDF extraction failed")
        raise ValueError("Could not read the PDF file") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))

        content = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text)

        return "\n".join(content)

    except Exception as exc:
        logger.exception("DOCX extraction failed")
        raise ValueError("Could not read the DOCX file") from exc


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore")

    except Exception as exc:
        logger.exception("TXT extraction failed")
        raise ValueError("Could not read the text file") from exc


# =========================================================
# RESUME STRENGTH
# =========================================================

def calculate_resume_strength(text: str) -> tuple[int, list[str]]:
    text_lower = text.lower()
    sections = find_sections(text)

    score = 0
    feedback = []

    # Contact information: 10 points
    email_found = bool(
        re.search(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
            text,
        )
    )

    phone_found = bool(
        re.search(
            r"(?:\+?\d[\d\s().-]{7,}\d)",
            text,
        )
    )

    if email_found:
        score += 5
    else:
        feedback.append("Add a professional email address")

    if phone_found:
        score += 5
    else:
        feedback.append("Add a contact phone number")

    # Core sections: 35 points
    core_sections = {
        "experience": 10,
        "education": 10,
        "skills": 10,
        "summary": 5,
    }

    section_feedback = {
        "experience": "Add a clear work experience section",
        "education": "Add a clear education section",
        "skills": "Add a dedicated skills section",
        "summary": "Add a short professional summary",
    }

    for section, points in core_sections.items():
        if sections.get(section):
            score += points
        else:
            feedback.append(section_feedback[section])

    # Projects: 10 points
    if sections.get("projects"):
        score += 10
    else:
        feedback.append(
            "Add a projects section to demonstrate practical experience"
        )

    # Relevant skills: 15 points
    detected_skills = find_known_skills(text)

    skill_count = len(detected_skills)

    if skill_count >= 8:
        score += 15
    elif skill_count >= 5:
        score += 11
        feedback.append(
            "Add a few more relevant technical or professional skills"
        )
    elif skill_count >= 2:
        score += 6
        feedback.append(
            "Expand the skills section with more role-relevant technologies"
        )
    else:
        feedback.append(
            "Add specific technical and professional skills"
        )

    # Action-oriented writing: 10 points
    action_verb_count = sum(
        1
        for verb in ACTION_VERBS
        if contains_term(text_lower, verb)
    )

    if action_verb_count >= 6:
        score += 10
    elif action_verb_count >= 3:
        score += 6
        feedback.append(
            "Use more strong action verbs to describe your work"
        )
    else:
        feedback.append(
            "Describe achievements with action verbs such as developed, built, improved, or implemented"
        )

    # Quantifiable achievements: 10 points
    measurable_patterns = re.findall(
        r"\b\d+(?:\.\d+)?\s?(?:%|percent|users?|clients?|projects?|hours?|days?|months?|years?|r|zar|\$|£|€)\b",
        text_lower,
    )

    if len(measurable_patterns) >= 3:
        score += 10
    elif len(measurable_patterns) >= 1:
        score += 6
        feedback.append(
            "Add more measurable achievements and results"
        )
    else:
        feedback.append(
            "Add measurable achievements using numbers, percentages, or results"
        )

    # Content depth: 10 points
    word_count = len(extract_words(text))

    if 300 <= word_count <= 1200:
        score += 10
    elif 180 <= word_count < 300:
        score += 6
        feedback.append(
            "Add more detail about your experience, projects, and achievements"
        )
    elif word_count > 1200:
        score += 6
        feedback.append(
            "Consider making the resume more concise and focused"
        )
    else:
        feedback.append(
            "The resume appears too brief; add more relevant detail"
        )

    return clamp_score(score), feedback


# =========================================================
# ATS COMPATIBILITY
# =========================================================

def calculate_ats_compatibility(text: str) -> tuple[int, list[str]]:
    sections = find_sections(text)

    score = 0
    feedback = []

    # Readable extracted text: 20 points
    word_count = len(extract_words(text))

    if word_count >= 300:
        score += 20
    elif word_count >= 180:
        score += 14
        feedback.append(
            "Add more detailed resume content for stronger ATS parsing"
        )
    else:
        score += 6
        feedback.append(
            "The resume contains very little readable text"
        )

    # Standard section headings: 30 points
    standard_sections = [
        "experience",
        "education",
        "skills",
    ]

    found_standard_sections = sum(
        1
        for section in standard_sections
        if sections.get(section)
    )

    score += found_standard_sections * 10

    if not sections.get("experience"):
        feedback.append(
            "Use a standard heading such as Work Experience or Professional Experience"
        )

    if not sections.get("education"):
        feedback.append(
            "Use a standard Education heading"
        )

    if not sections.get("skills"):
        feedback.append(
            "Use a standard Skills or Technical Skills heading"
        )

    # Contact details: 15 points
    email_found = bool(
        re.search(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
            text,
        )
    )

    phone_found = bool(
        re.search(
            r"(?:\+?\d[\d\s().-]{7,}\d)",
            text,
        )
    )

    if email_found:
        score += 8
    else:
        feedback.append(
            "Add a clearly readable email address"
        )

    if phone_found:
        score += 7
    else:
        feedback.append(
            "Add a clearly readable phone number"
        )

    # Skills and keywords: 20 points
    detected_skills = find_known_skills(text)
    skill_count = len(detected_skills)

    if skill_count >= 10:
        score += 20
    elif skill_count >= 6:
        score += 15
    elif skill_count >= 3:
        score += 9
        feedback.append(
            "Add more specific skills and technologies"
        )
    else:
        score += 3
        feedback.append(
            "The resume needs more searchable role-specific keywords"
        )

    # Bullet and achievement signals: 15 points
    bullet_count = len(
        re.findall(r"(?:^|\n)\s*[•●▪◦*-]\s+", text)
    )

    achievement_count = len(
        re.findall(
            r"\b\d+(?:\.\d+)?\s?(?:%|percent|users?|clients?|projects?|hours?|days?|months?|years?)\b",
            text.lower(),
        )
    )

    if bullet_count >= 5:
        score += 7
    elif bullet_count >= 2:
        score += 4
    else:
        feedback.append(
            "Use clear bullet points for experience and project achievements"
        )

    if achievement_count >= 3:
        score += 8
    elif achievement_count >= 1:
        score += 4
        feedback.append(
            "Add more quantified achievements"
        )
    else:
        feedback.append(
            "Include measurable results to strengthen ATS content"
        )

    return clamp_score(score), feedback


# =========================================================
# JOB MATCH
# =========================================================

def calculate_job_match(
    cv_text: str,
    job_text: str,
) -> tuple[Optional[int], list[str]]:

    job_text = normalise_text(job_text)

    if not job_text:
        return None, [
            "Paste a job description to calculate job fit"
        ]

    job_words = extract_words(job_text)

    if len(job_words) < 20:
        return None, [
            "Paste a fuller job description for a meaningful job match score"
        ]

    cv_keywords = extract_meaningful_keywords(cv_text)
    job_keywords = extract_meaningful_keywords(job_text)

    cv_skills = find_known_skills(cv_text)
    job_skills = find_known_skills(job_text)

    feedback = []

    # General keyword coverage: 45%
    if job_keywords:
        keyword_matches = cv_keywords.intersection(job_keywords)

        keyword_score = (
            len(keyword_matches) / len(job_keywords)
        ) * 100
    else:
        keyword_score = 0

    # Known skill coverage: 40%
    if job_skills:
        matched_skills = cv_skills.intersection(job_skills)

        skill_score = (
            len(matched_skills) / len(job_skills)
        ) * 100

        missing_skills = sorted(
            job_skills - cv_skills
        )

        if missing_skills:
            displayed_missing = ", ".join(
                missing_skills[:6]
            )

            feedback.append(
                f"Consider addressing these job-related skills if you genuinely have them: {displayed_missing}"
            )
    else:
        skill_score = keyword_score

    # Job-title / role terminology overlap: 15%
    role_terms = {
        word
        for word in job_keywords
        if len(word) >= 5
    }

    if role_terms:
        role_matches = role_terms.intersection(cv_keywords)

        role_score = (
            len(role_matches) / len(role_terms)
        ) * 100
    else:
        role_score = 0

    final_score = (
        keyword_score * 0.45
        + skill_score * 0.40
        + role_score * 0.15
    )

    final_score = clamp_score(final_score)

    if final_score >= 75:
        feedback.append(
            "Strong overall alignment with the job description"
        )
    elif final_score >= 50:
        feedback.append(
            "Moderate job alignment; tailor your resume more closely to the role"
        )
    else:
        feedback.append(
            "Low job alignment; focus the resume on relevant experience, skills, and achievements"
        )

    return final_score, feedback


# =========================================================
# COMBINED FEEDBACK
# =========================================================

def build_feedback(
    resume_feedback: list[str],
    ats_feedback: list[str],
    job_feedback: list[str],
) -> list[str]:

    combined = []

    for message in (
        resume_feedback
        + ats_feedback
        + job_feedback
    ):
        if message not in combined:
            combined.append(message)

    if not combined:
        combined.append(
            "Strong resume overall; continue tailoring it to each job description"
        )

    return combined[:8]


# =========================================================
# API ROUTE
# =========================================================

@router.post("/upload")
async def upload_cv(
    file: UploadFile = File(...),
    job_description: str = Form(""),
):
    try:
        if not file.filename:
            raise HTTPException(
                status_code=400,
                detail="No file name was provided",
            )

        safe_filename = Path(file.filename).name
        extension = Path(safe_filename).suffix.lower()

        if extension not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Use PDF, DOCX, or TXT.",
            )

        if (
            file.content_type
            and file.content_type not in ALLOWED_CONTENT_TYPES
        ):
            raise HTTPException(
                status_code=400,
                detail="Invalid file content type",
            )

        file_bytes = await file.read(
            MAX_FILE_SIZE + 1
        )

        if not file_bytes:
            raise HTTPException(
                status_code=400,
                detail="The uploaded file is empty",
            )

        if len(file_bytes) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail="File too large. Maximum size is 5MB.",
            )

        clean_job_description = normalise_text(
            job_description[:MAX_JOB_DESCRIPTION_LENGTH]
        )

        try:
            if extension == ".pdf":
                text = extract_text_from_pdf(file_bytes)

            elif extension == ".docx":
                text = extract_text_from_docx(file_bytes)

            else:
                text = extract_text_from_txt(file_bytes)

        except ValueError as exc:
            raise HTTPException(
                status_code=400,
                detail=str(exc),
            ) from exc

        text = normalise_text(text)

        if len(text) < MIN_EXTRACTED_TEXT_LENGTH:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Could not extract enough readable text "
                    "from the resume"
                ),
            )

        resume_score, resume_feedback = (
            calculate_resume_strength(text)
        )

        ats_score, ats_feedback = (
            calculate_ats_compatibility(text)
        )

        job_match_score, job_feedback = (
            calculate_job_match(
                text,
                clean_job_description,
            )
        )

        feedback = build_feedback(
            resume_feedback,
            ats_feedback,
            job_feedback,
        )

        return {
            "score": resume_score,
            "max_score": MAX_SCORE,
            "feedback": feedback,
            "ats_score": ats_score,
            "job_match_score": job_match_score,
        }

    except HTTPException:
        raise

    except Exception:
        logger.exception(
            "Unexpected error while analysing resume"
        )

        raise HTTPException(
            status_code=500,
            detail="Internal server error",
        )

    finally:
        await file.close()